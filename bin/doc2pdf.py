#!/usr/bin/env python3
"""Document to PDF Converter.

Converts Word documents (.doc/.docx/.wps) to PDF using LibreOffice (best formatting)
or pure Python fallback (olefile/python-docx text extraction).

Priority: LibreOffice > Pure Python (always works, no external deps).

Usage:
    python3 bin/doc2pdf.py input.docx -o output.pdf
    python3 bin/doc2pdf.py input.doc -o output.pdf
    python3 bin/doc2pdf.py document.wps -o output.pdf
    python3 bin/doc2pdf.py input.docx -o output.pdf --engine python
"""

import argparse
import os
import shutil
import subprocess
import sys


def check_command(cmd):
    return shutil.which(cmd) is not None


def convert_with_libreoffice(input_path, output_path):
    """Convert using LibreOffice headless mode (best formatting)."""
    output_dir = os.path.dirname(os.path.abspath(output_path))
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_path],
                   capture_output=True, timeout=120, check=True)
    expected = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
    if expected != os.path.abspath(output_path) and os.path.exists(expected):
        os.replace(expected, os.path.abspath(output_path))


def convert_docx_python(input_path, output_path):
    """Pure Python: python-docx -> PyMuPDF."""
    from docx import Document
    import fitz
    doc = Document(input_path)
    pdf = fitz.open()
    page = pdf.new_page()
    margin = 50
    y = margin
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y += 10; continue
        style = para.style.name if para.style else ""
        size = 11
        if 'Heading' in style or 'heading' in style:
            try: size = {1: 20, 2: 16, 3: 14}.get(int(style.replace('Heading ','').replace('heading ','')), 11)
            except: pass
        th = max(len(text) * size / 500 * 14 + 14, 14)
        if y + th > 820: page = pdf.new_page(); y = margin
        page.insert_textbox(fitz.Rect(margin, y, 545, y + th), text, fontsize=size, fontname="china-ss")
        y += th + 4
    for table in doc.tables:
        rows, cols = len(table.rows), len(table.columns)
        th = rows * 18 + 10
        if y + th > 820: page = pdf.new_page(); y = margin
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip()
                if ct:
                    r = fitz.Rect(margin + ci * 445 / cols, y, margin + (ci + 1) * 445 / cols, y + 16)
                    page.insert_textbox(r, ct, fontsize=9, fontname="china-ss")
        y += th + 10
    pdf.save(output_path)
    pdf.close()


def convert_doc_wps_python(input_path, output_path):
    """Pure Python: olefile text extraction -> PyMuPDF."""
    import fitz
    try: import olefile
    except ImportError: raise RuntimeError("olefile not found. Install: pip install olefile")
    try:
        ole = olefile.OleFileIO(input_path)
        data = ole.openstream('WordDocument').read()
        ole.close()
    except Exception as e: raise RuntimeError(f"Failed to read file: {e}")
    raw = data.decode('utf-16-le', errors='replace')
    safe = []
    for c in raw:
        cp = ord(c)
        if cp == 0: continue
        if cp in (10, 13): safe.append('\n'); continue
        if cp < 32: continue
        if 0x20 <= cp <= 0x7E or 0x4E00 <= cp <= 0x9FFF or 0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF:
            safe.append(c)
    text = ''.join(safe)
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        cjk = sum(1 for c in line if 0x4E00 <= ord(c) <= 0x9FFF)
        alpha = sum(1 for c in line if c.isascii() and c.isalpha())
        if cjk == 0 and alpha < 3: continue
        lines.append(line)
    if not lines: raise RuntimeError("Could not extract readable text")
    pdf = fitz.open()
    page = pdf.new_page()
    margin = 50; y = margin
    for line in lines:
        if y > 820: page = pdf.new_page(); y = margin
        x = margin
        for ch in line:
            cp = ord(ch)
            if cp < 32 or 0xD800 <= cp <= 0xDFFF: continue
            try:
                page.insert_text(fitz.Point(x, y), ch, fontsize=11, fontname="china-ss")
                x += 8
            except: pass
        y += 15
    pdf.save(output_path)
    pdf.close()


def convert(input_path, output_path, engine="auto"):
    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if engine in ("auto", "libreoffice"):
        if engine == "auto" or engine == "libreoffice":
            if check_command("soffice") or check_command("libreoffice"):
                try:
                    if engine == "auto": print("Using LibreOffice...")
                    convert_with_libreoffice(input_path, output_path)
                    print(f"Done: {output_path}")
                    return
                except Exception as e:
                    print(f"  LibreOffice failed: {e}", file=sys.stderr)
                    if engine == "libreoffice": sys.exit(1)

    if engine in ("auto", "python"):
        ext = os.path.splitext(input_path)[1].lower()
        print("Using pure Python conversion (text only, formatting may be lost)...")
        try:
            if ext == '.docx':
                convert_docx_python(input_path, output_path)
            else:
                convert_doc_wps_python(input_path, output_path)
            print(f"Done: {output_path}")
        except Exception as e:
            print(f"Error: {e}", file=sys.stderr)
            print(file=sys.stderr)
            print("Recommendation: Install LibreOffice for format-perfect conversion:", file=sys.stderr)
            if sys.platform == "darwin":
                print("  brew install --cask libreoffice", file=sys.stderr)
            elif sys.platform == "win32":
                print("  winget install TheDocumentFoundation.LibreOffice", file=sys.stderr)
            else:
                print("  sudo apt install libreoffice", file=sys.stderr)
            sys.exit(1)
    else:
        print(f"Error: Unknown engine '{engine}'", file=sys.stderr)
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser(
        description="Convert Word documents (.doc/.docx/.wps) to PDF.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s report.docx -o report.pdf
  %(prog)s report.doc -o report.pdf
  %(prog)s document.wps -o document.pdf
  %(prog)s input.docx -o output.pdf --engine python

Engines:
  libreoffice - LibreOffice headless (recommended, best formatting)
  python      - Pure Python (always works, text extraction, no formatting)
        """)
    parser.add_argument("input", help="Input document file (.doc/.docx/.wps)")
    parser.add_argument("-o", "--output", required=True, help="Output PDF file path")
    parser.add_argument("--engine", choices=["auto", "libreoffice", "python"],
                        default="auto", help="Conversion engine (default: auto, LibreOffice first)")
    args = parser.parse_args()

    convert(args.input, args.output, args.engine)


if __name__ == "__main__":
    main()