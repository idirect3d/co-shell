#!/usr/bin/env python3
"""Document to Markdown Converter.

Converts Word documents (.doc/.docx/.wps) to Markdown files,
preserving heading hierarchy, lists, tables, bold/italic formatting, and links.

Priority: python-docx (.docx) > LibreOffice fallback (.doc/.wps)

Usage:
    python3 bin/doc2md.py input.docx -o output.md
    python3 bin/doc2md.py input.doc -o output.md
    python3 bin/doc2md.py document.wps -o output.md
"""

import argparse
import os
import shutil
import subprocess
import sys
import tempfile


def check_command(cmd):
    return shutil.which(cmd) is not None


def export_html_with_libreoffice(input_path, output_path):
    """Export document to HTML using LibreOffice headless."""
    output_dir = os.path.dirname(os.path.abspath(output_path))
    subprocess.run(["soffice", "--headless", "--convert-to", "html",
                    "--outdir", output_dir, input_path],
                   capture_output=True, timeout=120, check=True)
    expected = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".html")
    if expected != os.path.abspath(output_path) and os.path.exists(expected):
        os.replace(expected, os.path.abspath(output_path))
    return True


def convert_docx_to_md(input_path, output_path):
    """Convert .docx file to Markdown using python-docx."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(input_path)
    md_lines = []

    # Heading style mapping
    heading_map = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        'Heading 4': 4,
        'Heading 5': 5,
        'Heading 6': 6,
        'Heading 7': 7,
        'Heading 8': 8,
        'Heading 9': 9,
    }

    # List tracking
    list_stack = []  # Stack of (level, marker) for nested lists

    def get_paragraph_style(paragraph):
        """Get the style name of a paragraph."""
        try:
            return paragraph.style.name if paragraph.style else ''
        except Exception:
            return ''

    def get_paragraph_alignment(paragraph):
        """Get paragraph alignment."""
        try:
            alignment = paragraph.alignment
            if alignment is None:
                return 'left'
            return alignment.value
        except Exception:
            return 'left'

    def format_run_text(run):
        """Format a run with bold/italic/underline."""
        text = run.text
        if not text:
            return ''

        # Escape markdown special characters
        text = text.replace('\\', '\\\\')
        text = text.replace('*', '\\*')
        text = text.replace('_', '\\_')
        text = text.replace('[', '\\[')
        text = text.replace(']', '\\]')
        text = text.replace('#', '\\#')
        text = text.replace('`', '\\`')
        text = text.replace('|', '\\|')

        if run.bold and run.italic:
            text = f'***{text}***'
        elif run.bold:
            text = f'**{text}**'
        elif run.italic:
            text = f'*{text}*'

        if run.underline:
            text = f'<u>{text}</u>'

        return text

    def get_list_marker(level):
        """Get appropriate list marker based on nesting level."""
        if level == 0:
            return '- '
        elif level == 1:
            return '  - '
        elif level == 2:
            return '    - '
        elif level == 3:
            return '      - '
        else:
            return '      ' + '  ' * (level - 3) + '- '

    def process_inline_elements(run_or_cell):
        """Process inline elements (runs) in a paragraph or table cell."""
        result = ''
        if hasattr(run_or_cell, 'runs'):
            for run in run_or_cell.runs:
                result += format_run_text(run)
        elif hasattr(run_or_cell, 'text'):
            # Table cell
            for paragraph in run_or_cell.paragraphs:
                for run in paragraph.runs:
                    result += format_run_text(run)
                result += '\n'
        return result

    def close_list_if_needed(current_level):
        """Close list items if we've gone back to a shallower level."""
        while list_stack and list_stack[-1] >= current_level:
            list_stack.pop()

    for para in doc.paragraphs:
        style_name = get_paragraph_style(para)
        text = para.text.strip()

        # Handle headings
        if style_name in heading_map:
            close_list_if_needed(999)  # Close any open lists
            level = heading_map[style_name]
            md_lines.append(f'\n{"#" * level} {text}\n')
            continue

        # Handle list items
        if style_name.startswith('List'):
            # Determine list level from style name
            try:
                level_str = style_name.split()[-1] if style_name.split()[-1].isdigit() else '1'
                level = int(level_str) - 1
            except (ValueError, IndexError):
                level = 0

            close_list_if_needed(level)
            marker = get_list_marker(level)
            md_lines.append(f'{marker}{text}')
            list_stack.append(level)
            continue

        # Handle empty paragraphs
        if not text:
            md_lines.append('')
            continue

        # Handle table of contents
        if style_name == 'TOC Heading':
            close_list_if_needed(999)
            md_lines.append(f'\n## {text}\n')
            continue

        # Handle normal paragraphs - check for alignment
        alignment = get_paragraph_alignment(para)
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            md_lines.append(f'\n<center>{text}</center>\n')
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            md_lines.append(f'\n<div align="right">{text}</div>\n')
        else:
            md_lines.append(text)

    # Close any remaining lists
    list_stack.clear()

    # Process tables
    for table in doc.tables:
        md_lines.append('')
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            md_line = '| ' + ' | '.join(cells) + ' |'
            md_lines.append(md_line)
            if row_idx == 0:
                separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                md_lines.append(separator)

    # Join and clean up
    md_content = '\n'.join(md_lines)

    # Clean up excessive blank lines
    while '\n\n\n' in md_content:
        md_content = md_content.replace('\n\n\n', '\n\n')

    md_content = md_content.strip() + '\n'

    # Write output
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    return md_content


def convert_doc_wps_to_md(input_path, output_path):
    """Convert .doc/.wps file to Markdown using LibreOffice HTML fallback."""
    if not check_command("soffice") and not check_command("libreoffice"):
        raise RuntimeError("LibreOffice not found. Install: brew install --cask libreoffice")

    # Export to HTML first
    tmp_dir = tempfile.mkdtemp()
    try:
        html_path = os.path.join(tmp_dir, 'temp_output.html')
        export_html_with_libreoffice(input_path, html_path)

        # Read HTML and convert to simple Markdown
        with open(html_path, 'r', encoding='utf-8', errors='replace') as f:
            html_content = f.read()

        # Simple HTML to Markdown conversion
        import re
        md_content = html_content

        # Convert headings
        md_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'\n# \1\n', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n## \1\n', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n### \1\n', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'\n#### \1\n', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<h5[^>]*>(.*?)</h5>', r'\n##### \1\n', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<h6[^>]*>(.*?)</h6>', r'\n###### \1\n', md_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert bold/italic
        md_content = re.sub(r'<b([^>]*)>(.*?)</b>', r'**\2**', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<strong([^>]*)>(.*?)</strong>', r'**\2**', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<i([^>]*)>(.*?)</i>', r'*\2*', md_content, flags=re.DOTALL | re.IGNORECASE)
        md_content = re.sub(r'<em([^>]*)>(.*?)</em>', r'*\2*', md_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert lists
        md_content = re.sub(r'<li([^>]*)>(.*?)</li>', r'\n- \2', md_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert paragraphs
        md_content = re.sub(r'<p[^>]*>(.*?)</p>', r'\n\1\n', md_content, flags=re.DOTALL | re.IGNORECASE)

        # Remove remaining HTML tags
        md_content = re.sub(r'<[^>]+>', '', md_content)

        # Decode HTML entities
        import html
        md_content = html.unescape(md_content)

        # Clean up
        while '\n\n\n' in md_content:
            md_content = md_content.replace('\n\n\n', '\n\n')
        md_content = md_content.strip() + '\n'

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

    finally:
        # Cleanup temp files
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


def convert(input_path, output_path):
    """Convert Word document to Markdown."""
    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()

    try:
        if ext == '.docx':
            print(f"Converting {input_path} to Markdown...")
            content = convert_docx_to_md(input_path, output_path)
        elif ext in ('.doc', '.wps'):
            print(f"Converting {input_path} to Markdown (via LibreOffice)...")
            convert_doc_wps_to_md(input_path, output_path)
            content = open(output_path, 'r', encoding='utf-8').read()
        else:
            print(f"Error: Unsupported format: {ext}", file=sys.stderr)
            print("Supported formats: .doc, .docx, .wps", file=sys.stderr)
            sys.exit(1)

        # Print summary
        lines = content.split('\n')
        headings = [l for l in lines if l.startswith('#')]
        print(f"\nConversion complete!")
        print(f"  Output: {output_path}")
        print(f"  Total lines: {len(lines)}")
        print(f"  Headings found: {len(headings)}")

        # Show heading structure
        if headings:
            print(f"\nDocument structure:")
            for h in headings[:20]:  # Show first 20 headings
                print(f"  {h}")
            if len(headings) > 20:
                print(f"  ... and {len(headings) - 20} more headings")

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser(
        description="Convert Word documents (.doc/.docx/.wps) to Markdown files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s report.docx -o output.md
  %(prog)s report.doc -o output.md
  %(prog)s document.wps -o output.md

Supported formats: .doc, .docx, .wps
For .doc/.wps files, LibreOffice is required for best results.
        """)
    parser.add_argument("input", help="Input Word document file (.doc/.docx/.wps)")
    parser.add_argument("-o", "--output", default=None, help="Output Markdown file (default: <input>.md)")
    args = parser.parse_args()

    if args.output is None:
        base_name = os.path.splitext(os.path.basename(args.input))[0]
        args.output = os.path.join(os.path.dirname(os.path.abspath(args.input)), base_name + ".md")

    convert(args.input, args.output)


if __name__ == "__main__":
    main()