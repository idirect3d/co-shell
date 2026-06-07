#!/usr/bin/env python3
"""DOCX to PDF Converter.

Converts .docx files to PDF with pure Python fallback (no external office suite required).
Also attempts to use WPS Office or LibreOffice if available for better formatting.

Priority: WPS Office > LibreOffice > Pure Python (always works, never depends on external apps).

Usage:
    python3 bin/docx2pdf.py input.docx -o output.pdf
"""

import argparse
import os
import shutil
import subprocess
import sys


def check_command(cmd):
    return shutil.which(cmd) is not None


def find_wps_macos():
    """Find WPS Office on macOS by bundle ID or common paths."""
    import subprocess
    try:
        r = subprocess.run(["mdfind", "kMDItemCFBundleIdentifier == 'com.kingsoft.wpsoffice*'"],
                           capture_output=True, text=True, timeout=5)
        paths = [p.strip() for p in r.stdout.split('\n') if p.strip()]
        for p in paths:
            exe = os.path.join(p, "Contents", "MacOS", "wps")
            if os.path.isfile(exe) and os.access(exe, os.X_OK):
                return exe
    except Exception:
        pass
    # Common paths
    for p in [
        "/Applications/WPS Office.app/Contents/MacOS/wps",
        os.path.expanduser("~/Applications/WPS Office.app/Contents/MacOS/wps"),
    ]:
        if os.path.isfile(p) and os.access(p, os.X_OK):
            return p
    return None


def find_wps_windows():
    import glob
    for pattern in [
        r"C:\Program Files\WPS Office\*\wps.exe",
        r"C:\Program Files (x86)\WPS Office\*\wps.exe",
    ]:
        matches = glob.glob(pattern)
        if matches:
            return matches[0]
    return None


def convert_with_wps(input_path, output_path):
    """WPS Office conversion (Linux wps2pdf or macOS/Windows via automation)."""
    wps2pdf = shutil.which("wps2pdf")
    if wps2pdf:
        subprocess.run([wps2pdf, input_path, output_path], capture_output=True, timeout=120, check=True)
        return
    if sys.platform == "darwin":
        wps = find_wps_macos()
        if wps:
            subprocess.run([wps, os.path.abspath(input_path), '--export-to-pdf', os.path.abspath(output_path)],
                           capture_output=True, timeout=120, shell=True)
            return
    if sys.platform == "win32":
        wps = find_wps_windows()
        if wps:
            try:
                import win32com.client
                app = win32com.client.Dispatch("KWps.Application")
                app.Visible = False
                doc = app.Documents.Open(os.path.abspath(input_path))
                doc.ExportAsFixedFormat(os.path.abspath(output_path), 17)
                doc.Close()
                app.Quit()
                return
            except ImportError:
                pass
    raise RuntimeError("WPS Office not found or conversion failed")


def convert_with_libreoffice(input_path, output_path):
    """LibreOffice headless conversion."""
    output_dir = os.path.dirname(os.path.abspath(output_path))
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_path],
                   capture_output=True, timeout=120, check=True)
    expected = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
    if expected != os.path.abspath(output_path) and os.path.exists(expected):
        os.replace(expected, os.path.abspath(output_path))


def convert_with_python(input_path, output_path):
    """Pure Python conversion: python-docx -> PyMuPDF (never requires external apps)."""
    from docx import Document
    import fitz

    doc = Document(input_path)
    pdf = fitz.open()
    page = None
    y = 40  # current y position on page
    margin = 50
    page_width = 595  # A4 at 72 DPI
    page_height = 842
    line_height = 14
    font_size = 11

    def new_page():
        nonlocal page, y
        page = pdf.new_page(width=page_width, height=page_height)
        y = margin

    new_page()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y += line_height
            continue

        # Determine font size based on style
        style_name = para.style.name if para.style else ""
        if 'Heading' in style_name or 'heading' in style_name:
            try:
                level = int(style_name.replace('Heading ', '').replace('heading ', ''))
            except ValueError:
                level = 1
            size = {1: 20, 2: 16, 3: 14, 4: 12}.get(level, 11)
            bold = True
        else:
            size = font_size
            bold = False

        # Check page space
        text_height = max(len(text) * size / page_width * line_height + line_height, line_height)
        if y + text_height > page_height - margin:
            new_page()

        # Draw text (use china-ss for CJK support)
        font = "china-ss" if not bold else "china-ss"
        rect = fitz.Rect(margin, y, page_width - margin, y + text_height)
        page.insert_textbox(rect, text, fontsize=size, fontname="china-ss", color=(0, 0, 0))
        if bold:
            # Simulate bold by drawing a second copy offset slightly
            rect2 = fitz.Rect(margin + 0.5, y, page_width - margin + 0.5, y + text_height)
            page.insert_textbox(rect2, text, fontsize=size, fontname="china-ss", color=(0, 0, 0))
        y += text_height + 4

    # Handle tables
    for table in doc.tables:
        rows, cols = len(table.rows), len(table.columns)
        table_height = rows * 20 + 10
        if y + table_height > page_height - margin:
            new_page()
        x = margin
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    cell_rect = fitz.Rect(x + ci * (page_width - 2 * margin) / cols, y,
                                          x + (ci + 1) * (page_width - 2 * margin) / cols, y + 18)
                    page.insert_textbox(cell_rect, cell_text, fontsize=9, fontname="china-ss", color=(0, 0, 0))
        y += table_height + 10

    pdf.save(output_path)
    pdf.close()


def convert(input_path, output_path, engine="auto"):
    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if engine == "auto":
        # 1. Try WPS
        wps_ok = False
        if sys.platform == "linux" and check_command("wps2pdf"):
            wps_ok = True
        elif sys.platform == "darwin" and find_wps_macos():
            wps_ok = True
        elif sys.platform == "win32" and find_wps_windows():
            wps_ok = True

        if wps_ok:
            try:
                print("Trying WPS Office...")
                convert_with_wps(input_path, output_path)
                print(f"Done: {output_path}")
                return
            except Exception as e:
                print(f"  WPS failed: {e}", file=sys.stderr)

        # 2. Try LibreOffice
        if check_command("soffice") or check_command("libreoffice"):
            try:
                print("Trying LibreOffice...")
                convert_with_libreoffice(input_path, output_path)
                print(f"Done: {output_path}")
                return
            except Exception as e:
                print(f"  LibreOffice failed: {e}", file=sys.stderr)

        # 3. Pure Python (always works)
        print("Using pure Python conversion...")
        convert_with_python(input_path, output_path)
        print(f"Done: {output_path}")

    elif engine == "wps":
        convert_with_wps(input_path, output_path)
        print(f"Done: {output_path}")
    elif engine == "libreoffice":
        convert_with_libreoffice(input_path, output_path)
        print(f"Done: {output_path}")
    elif engine == "python":
        convert_with_python(input_path, output_path)
        print(f"Done: {output_path}")
    else:
        print(f"Error: Unknown engine '{engine}'", file=sys.stderr)
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser(
        description="Convert DOCX files to PDF format.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s report.docx -o report.pdf
  %(prog)s report.docx -o report.pdf --engine wps
  %(prog)s report.docx -o report.pdf --engine python

Engines (in order for auto mode):
  wps         - WPS Office (if installed)
  libreoffice - LibreOffice (if installed)
  python      - Pure Python (always works, no external deps)
        """)
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("-o", "--output", required=True, help="Output PDF file path")
    parser.add_argument("--engine", choices=["auto", "wps", "libreoffice", "python"],
                        default="auto", help="Conversion engine (default: auto-detect)")
    args = parser.parse_args()

    convert(args.input, args.output, args.engine)


if __name__ == "__main__":
    main()