#!/usr/bin/env python3
"""Markdown to Beautiful Word Document Converter.

Supports multiple styles including official Chinese government document format
(GB/T 9704-2012). Default style is 'official'.

Usage:
    python3 md2docx.py input.md -o output.docx
    python3 md2docx.py input.md --style modern
    python3 md2docx.py input.md --style official
"""

import argparse
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _rgb(r, g, b):
    """Create RGBColor."""
    return RGBColor(r, g, b)


STYLES = {
    "modern": {
        "primary": _rgb(0x1A, 0x73, 0xE8),
        "secondary": _rgb(0x5F, 0x63, 0x68),
        "heading1": _rgb(0x1A, 0x73, 0xE8),
        "heading2": _rgb(0x19, 0x60, 0x7D),
        "heading3": _rgb(0x37, 0x47, 0x4F),
        "body": _rgb(0x33, 0x33, 0x33),
        "code_bg": "E8F0FE",
        "code_fg": _rgb(0x1D, 0x1D, 0x1D),
        "link": _rgb(0x1A, 0x0D, 0xAB),
        "blockquote_bg": "E8EAF6",
        "table_header": "1A73E8",
        "table_alt_row": "F1F3F4",
    },
    "classic": {
        "primary": _rgb(0x2C, 0x3E, 0x50),
        "secondary": _rgb(0x7F, 0x8C, 0x8D),
        "heading1": _rgb(0x2C, 0x3E, 0x50),
        "heading2": _rgb(0x2C, 0x3E, 0x50),
        "heading3": _rgb(0x7F, 0x8C, 0x8D),
        "body": _rgb(0x1A, 0x1A, 0x1A),
        "code_bg": "ECF0F1",
        "code_fg": _rgb(0x2C, 0x3E, 0x50),
        "link": _rgb(0x29, 0x80, 0xB9),
        "blockquote_bg": "F5F5F5",
        "table_header": "2C3E50",
        "table_alt_row": "F8F9FA",
    },
    "minimal": {
        "primary": _rgb(0x00, 0x00, 0x00),
        "secondary": _rgb(0x55, 0x55, 0x55),
        "heading1": _rgb(0x00, 0x00, 0x00),
        "heading2": _rgb(0x33, 0x33, 0x33),
        "heading3": _rgb(0x55, 0x55, 0x55),
        "body": _rgb(0x33, 0x33, 0x33),
        "code_bg": "F5F5F5",
        "code_fg": _rgb(0x33, 0x33, 0x33),
        "link": _rgb(0x55, 0x55, 0x55),
        "blockquote_bg": "FAFAFA",
        "table_header": "333333",
        "table_alt_row": "F5F5F5",
    },
    "official": {
        # GB/T 9704-2012 党政机关公文格式
        "primary": _rgb(0x00, 0x00, 0x00),
        "secondary": _rgb(0x00, 0x00, 0x00),
        "heading1": _rgb(0x00, 0x00, 0x00),
        "heading2": _rgb(0x00, 0x00, 0x00),
        "heading3": _rgb(0x00, 0x00, 0x00),
        "body": _rgb(0x00, 0x00, 0x00),
        "code_bg": "F5F5F5",
        "code_fg": _rgb(0x00, 0x00, 0x00),
        "link": _rgb(0x00, 0x00, 0x00),
        "blockquote_bg": "F5F5F5",
        "table_header": "000000",
        "table_alt_row": "F5F5F5",
    },
}

HEADING_SIZES = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}


def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    shading = parse_xml('<w:shd ' + nsdecls('w') + ' w:fill="' + color_hex + '"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_spacing(para, before=0, after=0, line=None, line_rule=None):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        if line_rule == WD_LINE_SPACING.EXACTLY:
            pf.line_spacing = Pt(line)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        else:
            pf.line_spacing = line


def set_run_font(run, font_name=None, font_east_asia=None):
    """Set font for a run, including East-Asian font."""
    if font_name:
        run.font.name = font_name
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml('<w:rPr ' + nsdecls('w') + '/>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts ' + nsdecls('w') + '/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        if font_east_asia:
            rFonts.set(qn('w:eastAsia'), font_east_asia)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)


def add_run(paragraph, text, bold=False, italic=False, color=None,
            size=None, font_name=None, font_east_asia=None, underline=False):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font_name:
        set_run_font(run, font_name, font_east_asia)
    if underline:
        run.font.underline = True
    return run


def color_to_hex(c):
    """Convert RGBColor to hex string."""
    return '%02X%02X%02X' % (c[0], c[1], c[2])


def add_bottom_border(para, color_hex, sz=6):
    """Add bottom border to paragraph."""
    text = '<w:pBdr ' + nsdecls('w') + '>'
    text += '  <w:bottom w:val="single" w:sz="' + str(sz) + '" w:space="4" w:color="' + color_hex + '"/>'
    text += '</w:pBdr>'
    para._p.get_or_add_pPr().append(parse_xml(text))


def add_left_border(para, color_hex, sz=18):
    """Add left border to paragraph."""
    text = '<w:pBdr ' + nsdecls('w') + '>'
    text += '  <w:left w:val="single" w:sz="' + str(sz) + '" w:space="8" w:color="' + color_hex + '"/>'
    text += '</w:pBdr>'
    para._p.get_or_add_pPr().append(parse_xml(text))


def add_shading(para, color_hex):
    """Add background shading."""
    text = '<w:shd ' + nsdecls('w') + ' w:fill="' + color_hex + '" w:val="clear"/>'
    para._p.get_or_add_pPr().append(parse_xml(text))


def add_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr ' + nsdecls('w') + '/>')
        tbl.insert(0, tblPr)
    text = '<w:tblBorders ' + nsdecls('w') + '>'
    text += '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    text += '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    text += '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    text += '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    text += '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    text += '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    text += '</w:tblBorders>'
    tblPr.append(parse_xml(text))


def parse_inline(text):
    """Parse markdown inline formatting.

    Supports: ***bold+italic***, **bold**, *italic*, `code`, [link](url)
    Returns list of (text, properties_dict).
    """
    tokens = []
    # Order matters: *** must be checked before ** and *
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))'
    last_end = 0
    for m in re.finditer(pattern, text):
        start, end = m.start(), m.end()
        if start > last_end:
            tokens.append((text[last_end:start], {}))
        full = m.group(0)
        if full.startswith('***'):
            tokens.append((m.group(2), {'bold': True, 'italic': True}))
        elif full.startswith('**'):
            tokens.append((m.group(3), {'bold': True}))
        elif full.startswith('*'):
            tokens.append((m.group(4), {'italic': True}))
        elif full.startswith('`'):
            tokens.append((m.group(5), {'code': True}))
        elif full.startswith('['):
            tokens.append((m.group(6), {'link': m.group(7)}))
        last_end = end
    if last_end < len(text):
        tokens.append((text[last_end:], {}))
    return tokens if tokens else [(text, {})]


def apply_inline(paragraph, text, style_name, font_size=11, font_family='Arial',
                 font_east_asia=None):
    """Apply inline markdown formatting to a paragraph."""
    s = STYLES[style_name]
    for token_text, props in parse_inline(text):
        if not token_text:
            continue
        run = paragraph.add_run(token_text)
        run.font.name = font_family
        run.font.size = Pt(font_size)
        set_run_font(run, font_family, font_east_asia)

        if props.get('bold'):
            run.bold = True
        if props.get('italic'):
            run.italic = True
        if props.get('code'):
            run.font.name = 'Courier New'
            run.font.color.rgb = s['code_fg']
        elif props.get('link'):
            run.font.color.rgb = s['link']
            run.font.underline = True
        else:
            run.font.color.rgb = s['body']


def build_cover_page(doc, title, author, style_name):
    """Build a cover page."""
    s = STYLES[style_name]
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '\u2501' * 40, color=s['primary'], size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, bold=True, color=s['primary'], size=28, font_name='Arial')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '\u2501' * 40, color=s['primary'], size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'By ' + author, color=s['secondary'], size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, datetime.now().strftime('%B %d, %Y'), color=s['secondary'], size=11)
    doc.add_page_break()


def build_toc(doc, headings, style_name):
    """Build a table of contents."""
    s = STYLES[style_name]
    p = doc.add_paragraph()
    add_run(p, 'Table of Contents', bold=True, color=s['primary'], size=18)
    set_spacing(p, after=12)
    p = doc.add_paragraph()
    add_run(p, '\u2501' * 60, color=s['secondary'], size=8)
    set_spacing(p, after=6)
    for level, text in headings:
        indent = '    ' * (level - 1)
        bullets = {1: '\u25CF', 2: '\u25CB', 3: '\u25AA'}
        bullet = bullets.get(level, '\u2022')
        sz = 12 - level if level <= 2 else 10
        p = doc.add_paragraph()
        add_run(p, indent + bullet + '  ' + text,
                color=s['body'] if level <= 3 else s['secondary'],
                size=max(sz, 9))
        set_spacing(p, before=2, after=2)
    doc.add_page_break()


def build_heading(doc, text, level, style_name, base_size):
    """Build a styled heading."""
    s = STYLES[style_name]
    colors = {1: s['heading1'], 2: s['heading2'], 3: s['heading3']}
    color = colors.get(level, s['body'])
    sz = HEADING_SIZES.get(level, base_size)
    p = doc.add_paragraph()
    add_run(p, text, bold=True, color=color, size=sz, font_name='Arial')
    if level <= 2:
        add_bottom_border(p, color_to_hex(color))
    sb = {1: 24, 2: 18, 3: 12, 4: 8, 5: 6, 6: 6}
    sa = {1: 12, 2: 8, 3: 6, 4: 4, 5: 3, 6: 3}
    set_spacing(p, before=sb.get(level, 6), after=sa.get(level, 4))


def build_heading_official(doc, text, doc_level):
    """Build a heading in official Chinese government document format.

    GB/T 9704-2012:
    - doc_level=0 (主标题/文档标题)：小标宋, 二号 (22pt), centered, 不加粗
    - doc_level=1 (一级标题，如一、)：黑体, 三号 (16pt), 不加粗, 首行缩进2字符
    - doc_level=2 (二级标题，如（一）)：楷体_GB2312, 三号 (16pt), 加粗, 首行缩进2字符
    - doc_level=3+ (三级及以下)：仿宋_GB2312, 三号 (16pt), 加粗, 首行缩进2字符

    Supports inline formatting (bold/italic/code/link) within heading text.
    """
    if doc_level == 0:
        # Document title (主标题): 小标宋, 二号 (22pt), centered, 不加粗
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_inline(p, text, 'official', font_size=22, font_family='小标宋',
                     font_east_asia='小标宋')
        set_spacing(p, before=0, after=0, line=28, line_rule=WD_LINE_SPACING.EXACTLY)
    elif doc_level == 1:
        # Level 1 heading (一级标题): 黑体, 三号 (16pt), 不加粗, 首行缩进2字符
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Pt(16 * 2)  # 2 character indent
        apply_inline(p, text, 'official', font_size=16, font_family='黑体',
                     font_east_asia='黑体')
        set_spacing(p, before=0, after=0, line=28, line_rule=WD_LINE_SPACING.EXACTLY)
    elif doc_level == 2:
        # Level 2 heading (二级标题): 楷体_GB2312, 三号 (16pt), 加粗, 首行缩进2字符
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Pt(16 * 2)  # 2 character indent
        apply_inline(p, text, 'official', font_size=16, font_family='楷体',
                     font_east_asia='楷体')
        for run in p.runs:
            run.bold = True
        set_spacing(p, before=0, after=0, line=28, line_rule=WD_LINE_SPACING.EXACTLY)
    else:
        # Level 3+ headings (三级及以下): 仿宋_GB2312, 三号 (16pt), 加粗, 首行缩进2字符
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Pt(16 * 2)  # 2 character indent
        apply_inline(p, text, 'official', font_size=16, font_family='仿宋',
                     font_east_asia='仿宋')
        for run in p.runs:
            run.bold = True
        set_spacing(p, before=0, after=0, line=28, line_rule=WD_LINE_SPACING.EXACTLY)


def build_code_block(doc, code_text, style_name, font_size):
    """Build a styled code block."""
    s = STYLES[style_name]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, s['code_bg'])
    cell.paragraphs[0].clear()
    for cl in code_text.split('\n'):
        p = cell.add_paragraph()
        run = p.add_run(cl if cl else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(font_size)
        run.font.color.rgb = s['code_fg']
        set_spacing(p, before=0, after=0, line=1.2)
    if cell.paragraphs[0].text.strip() == '':
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    doc.add_paragraph()


def build_blockquote(doc, text, style_name, font_size):
    """Build a styled blockquote."""
    s = STYLES[style_name]
    p = doc.add_paragraph()
    apply_inline(p, text, style_name, font_size)
    add_left_border(p, color_to_hex(s['primary']))
    p.paragraph_format.left_indent = Cm(1.0)
    add_shading(p, s['blockquote_bg'])
    set_spacing(p, before=6, after=6)


def build_list(doc, items, ordered, style_name, font_size):
    """Build a styled list."""
    s = STYLES[style_name]
    for idx, item_text in enumerate(items):
        p = doc.add_paragraph()
        prefix = str(idx + 1) + '. ' if ordered else '\u2022 '
        run = p.add_run(prefix)
        run.bold = True
        run.font.color.rgb = s['primary']
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        apply_inline(p, item_text, style_name, font_size)
        pf = p.paragraph_format
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-0.5)
        set_spacing(p, before=2, after=2)


def build_list_official(doc, items, ordered, font_size=16):
    """Build a list in official Chinese government document format.

    Uses 仿宋_GB2312, 三号 (16pt), 固定行距28磅.
    Supports inline formatting (bold/italic/code/link) within list items.
    """
    for idx, item_text in enumerate(items):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prefix = str(idx + 1) + '.\u3000' if ordered else '\u3000\u3000'
        run = p.add_run(prefix)
        run.font.name = '仿宋'
        run.font.size = Pt(font_size)
        set_run_font(run, '仿宋', '仿宋')

        # Add item text with inline formatting support
        apply_inline(p, item_text, 'official', font_size=font_size,
                     font_family='仿宋', font_east_asia='仿宋')

        pf = p.paragraph_format
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = Cm(-0.74)
        set_spacing(p, before=0, after=0, line=28, line_rule=WD_LINE_SPACING.EXACTLY)


def build_table(doc, data, style_name, font_size):
    """Build a styled table."""
    s = STYLES[style_name]
    if not data or not data[0]:
        return
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, ct in enumerate(data[0]):
        cell = table.cell(0, j)
        set_cell_shading(cell, s['table_header'])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ct)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
    for i in range(1, rows):
        for j in range(cols):
            cell = table.cell(i, j)
            if i % 2 == 0:
                set_cell_shading(cell, s['table_alt_row'])
            text = data[i][j] if j < len(data[i]) else ''
            p = cell.paragraphs[0]
            apply_inline(p, text, style_name, max(font_size - 1, 9))
    add_table_borders(table)
    doc.add_paragraph()


def build_image(doc, img_path, alt_text, style_name, input_dir):
    """Build an image element."""
    s = STYLES[style_name]
    if not os.path.isabs(img_path):
        img_path = os.path.join(input_dir, img_path)
    if os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(4.5))
            set_spacing(p, before=6, after=6)
            if alt_text:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(cap, alt_text, italic=True, color=s['secondary'], size=9)
                set_spacing(cap, after=6)
        except Exception as e:
            p = doc.add_paragraph()
            msg = '[Image: ' + alt_text + ' - error: ' + str(e) + ']'
            add_run(p, msg, color=s['primary'], size=10, italic=True)
    else:
        p = doc.add_paragraph()
        msg = '[Image: ' + alt_text + ' - not found: ' + img_path + ']'
        add_run(p, msg, color=s['primary'], size=10, italic=True)


def setup_official_page(doc):
    """Set up page layout for official Chinese government document.

    GB/T 9704-2012:
    - Paper: A4 (210mm x 297mm)
    - Top margin: 37mm
    - Bottom margin: 35mm
    - Left margin: 28mm
    - Right margin: 26mm
    """
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)


def build_paragraph_official(doc, text, font_size=16):
    """Build a body paragraph in official Chinese government document format.

    GB/T 9704-2012:
    - Font: 仿宋_GB2312, 三号 (16pt)
    - Line spacing: fixed 28pt
    - First line indent: 2 characters (2 x 16pt = 32pt ≈ 1.13cm)
    - Text alignment: justify (两端对齐)

    Supports inline formatting (bold/italic/code/link) within paragraph text.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Pt(font_size * 2)  # 2 character indent
    set_spacing(p, before=0, after=0, line=28, line_rule=WD_LINE_SPACING.EXACTLY)

    # Apply inline formatting (bold/italic/code/link)
    apply_inline(p, text, 'official', font_size=font_size,
                 font_family='仿宋', font_east_asia='仿宋')

    return p


# Chinese ordinal number patterns for heading level detection
# Level 1: 一、二、三、... (single Chinese number + 、)
# Level 2: （一）（二）（三）... (Chinese number in brackets)
# Level 3: 1. 2. 3. ... (Arabic number + .)
CN_ORDINAL_PATTERNS = {
    1: re.compile(r'^[一二三四五六七八九十百千]+[、]'),
    2: re.compile(r'^（[一二三四五六七八九十百千]+）'),
    3: re.compile(r'^\d+[\.\、]'),
}


def compute_official_heading_levels(headings):
    """Compute document heading levels for official style.

    Rules:
    - The first # heading is treated as the document title (doc_level=0)
    - Subsequent headings have their level shifted down by 1:
      ## -> doc_level=1 (一级标题), ### -> doc_level=2 (二级标题), etc.
    - If there is no # heading, the first ## is treated as doc_level=1

    Returns a dict mapping (raw_level, index) -> doc_level.
    """
    result = {}
    first_heading_seen = False
    for idx, (raw_level, text) in enumerate(headings):
        if not first_heading_seen:
            first_heading_seen = True
            if raw_level == 1:
                # First heading is # -> document title
                result[(raw_level, idx)] = 0
                continue
            else:
                # First heading is ## or lower -> no document title
                # The first heading becomes doc_level=1
                result[(raw_level, idx)] = 1
                continue
        # Subsequent headings: shift down by 1
        doc_level = raw_level - 1
        if doc_level < 1:
            doc_level = 1
        result[(raw_level, idx)] = doc_level
    return result


def convert(md_content, output_path, args):
    """Convert markdown to a styled Word document."""
    style_name = args.style
    base_size = args.font_size
    base_font = args.font_family
    s = STYLES[style_name]
    doc = Document()

    if style_name == 'official':
        setup_official_page(doc)
    else:
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Collect all headings for official style level computation
    headings = []
    for line in md_content.split('\n'):
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            headings.append((len(m.group(1)), m.group(2).strip()))

    # Compute heading doc_level mapping for official style
    heading_level_map = {}
    if style_name == 'official':
        heading_level_map = compute_official_heading_levels(headings)

    if not args.no_cover and style_name != 'official':
        build_cover_page(doc, args.title, args.author, style_name)
    if not args.no_toc and headings and style_name != 'official':
        build_toc(doc, headings, style_name)

    lines = md_content.split('\n')
    i = 0
    heading_idx = 0  # Track heading index for official style level lookup
    in_code_block = False
    code_buffer = []
    input_dir = os.path.dirname(os.path.abspath(args.input))

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith('```'):
            if in_code_block:
                build_code_block(doc, '\n'.join(code_buffer), style_name, base_size - 1)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and stripped.endswith('|'):
            table_data = []
            while i < len(lines) and '|' in lines[i].strip():
                cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                table_data.append(cells)
                i += 1
            if len(table_data) >= 2 and any('---' in c for c in table_data[1]):
                table_data.pop(1)
            if table_data:
                build_table(doc, table_data, style_name, base_size)
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            p = doc.add_paragraph()
            add_run(p, '\u2501' * 60, color=s['secondary'], size=8)
            set_spacing(p, before=6, after=6)
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            if style_name == 'official':
                raw_level = len(m.group(1))
                doc_level = heading_level_map.get((raw_level, heading_idx), 1)
                build_heading_official(doc, m.group(2).strip(), doc_level)
                heading_idx += 1
            else:
                build_heading(doc, m.group(2).strip(), len(m.group(1)),
                              style_name, base_size)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quotes = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quotes.append(lines[i].strip().lstrip('>').strip())
                i += 1
            if style_name == 'official':
                build_paragraph_official(doc, ' '.join(quotes))
            else:
                build_blockquote(doc, ' '.join(quotes), style_name, base_size)
            continue

        # Unordered list
        if re.match(r'^[\s]*[-*+]\s+', stripped):
            items = []
            while i < len(lines):
                sl = lines[i].strip()
                if re.match(r'^[\s]*[-*+]\s+', sl):
                    items.append(re.sub(r'^[\s]*[-*+]\s+', '', sl))
                    i += 1
                else:
                    break
            if style_name == 'official':
                build_list_official(doc, items, False)
            else:
                build_list(doc, items, False, style_name, base_size)
            continue

        # Ordered list
        if re.match(r'^\s*\d+[\.)]\s+', stripped):
            items = []
            while i < len(lines):
                sl = lines[i].strip()
                if re.match(r'^\s*\d+[\.)]\s+', sl):
                    items.append(re.sub(r'^\s*\d+[\.)]\s+', '', sl))
                    i += 1
                else:
                    break
            if style_name == 'official':
                build_list_official(doc, items, True)
            else:
                build_list(doc, items, True, style_name, base_size)
            continue

        # Image
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            build_image(doc, img_match.group(2), img_match.group(1),
                        style_name, input_dir)
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Regular paragraph
        if style_name == 'official':
            build_paragraph_official(doc, stripped, base_size)
        else:
            p = doc.add_paragraph()
            apply_inline(p, stripped, style_name, base_size, base_font)
            set_spacing(p, before=3, after=3)
        i += 1
    doc.save(output_path)
    return True


def main():
    """Entry point."""
    parser = argparse.ArgumentParser(
        description='Convert Markdown to a beautiful Word document.',
        formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('-o', '--output', help='Output Word file path')
    parser.add_argument('--style', choices=['modern', 'classic', 'minimal', 'official'],
                        default='official', help='Visual style (default: official/GB/T 9704-2012)')
    parser.add_argument('--title', help='Document title')
    parser.add_argument('--author', default='co-shell', help='Author name')
    parser.add_argument('--font-size', type=int, default=16,
                        help='Base font size in pt (default: 16 for official, 11 for others)')
    parser.add_argument('--font-family', default='仿宋', help='Font family (default: 仿宋 for official)')
    parser.add_argument('--no-toc', action='store_true',
                        help='Skip table of contents')
    parser.add_argument('--no-cover', action='store_true',
                        help='Skip cover page')
    parser.add_argument('--debug', action='store_true',
                        help='Print debug info')
    args = parser.parse_args()
    if not os.path.exists(args.input):
        print('Error: File not found: ' + args.input)
        sys.exit(1)
    if not args.output:
        base = os.path.splitext(os.path.basename(args.input))[0]
        args.output = base + '.docx'
    if not args.title:
        args.title = os.path.splitext(os.path.basename(args.input))[0]
    with open(args.input, 'r', encoding='utf-8') as f:
        md_content = f.read()
    if args.debug:
        print('Input:   ' + args.input)
        print('Output:  ' + args.output)
        print('Style:   ' + args.style)
        print('Font:    ' + args.font_family + ' ' + str(args.font_size) + 'pt')
        print('Title:   ' + args.title)
        print('Author:  ' + args.author)
        print('Size:    ' + str(len(md_content)) + ' chars')
        print('-' * 50)
    try:
        convert(md_content, args.output, args)
        size_kb = os.path.getsize(args.output) / 1024.0
        print('Successfully converted to: ' + args.output)
        print('File size: %.1f KB' % size_kb)
    except Exception as e:
        print('Error: ' + str(e))
        if args.debug:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
