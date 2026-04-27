#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
鸿运承物业管理中心2026年一季度"接诉即办"工作分析讲评
按照 GB/T 9704-2012 公文标准格式
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21.0)
s.page_height = Cm(29.7)
s.top_margin = Cm(3.7)
s.bottom_margin = Cm(3.5)
s.left_margin = Cm(2.8)
s.right_margin = Cm(2.6)

def ap(text, sz=16, fn='仿宋', b=False, c=False, ind=True):
    p = doc.add_paragraph()
    if c:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.name = fn
    r.bold = b
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    if ind and not c:
        pf.first_line_indent = Pt(32)

def add_complex_table():
    """创建按派单来源分列的对比表（严格参照2025年台账格式）"""
    # 表头结构：性质 | 12345(本期|同期) | 12328(本期|同期) | 政风行风(本期|同期) | 舆情(本期|同期) | 合计(本期|同期)
    headers_top = ['性质',
                   '12345', '', '12328', '', '风行风', '', '舆情', '',
                   '合计', '']
    headers_sub = ['', '本期', '同期', '本期', '同期', '本期', '同期', '本期', '同期',
                   '本期', '同期']

    # 数据行
    data_rows = [
        ['诉求', '182', '183', '0', '8', '0', '1', '', '3', '182', '195'],
        ['表扬', '13', '10', '0', '0', '0', '0', '0', '0', '13', '10'],
        ['其他', '50', '40', '0', '3', '0', '0', '0', '1', '50', '44'],
        ['咨询', '1', '2', '0', '0', '0', '0', '0', '0', '1', '2'],
        ['合计', '246', '235', '0', '11', '0', '1', '0', '4', '246', '251'],
    ]

    # 同比变化行
    change_rows = [
        ['同比', '—', '—', '—', '—', '—', '—', '—', '—', '—', '—'],
        ['', '减少39件', '增加--件', '减少1件', '增加4件', '减少5件'],
    ]

    ncols = len(headers_top)  # 11列
    nrows = 1 + 1 + len(data_rows)  # 表头行 + 副表头行 + 数据行

    t = doc.add_table(rows=nrows, cols=ncols)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置表格边框
    tblPr = t._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr ' + nsdecls('w') + '/>')
        t._tbl.insert(0, tblPr)
    borders = parse_xml(
        '<w:tblBorders ' + nsdecls('w') + '>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>')
    tblPr.append(borders)

    # 填充第0行（主表头）
    for i, h in enumerate(headers_top):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.size = Pt(10)
        r.font.name = '黑体'
        r.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 合并主表头的单元格（每个渠道占2列，性质占1列）
    # 合并 性质 列 (col 0)
    t.rows[0].cells[0].merge(t.rows[1].cells[0])
    # 合并 12345 (cols 1-2)
    t.rows[0].cells[1].merge(t.rows[0].cells[2])
    # 合并 12328 (cols 3-4)
    t.rows[0].cells[3].merge(t.rows[0].cells[4])
    # 合并 政风行风 (cols 5-6)
    t.rows[0].cells[5].merge(t.rows[0].cells[6])
    # 合并 舆情 (cols 7-8)
    t.rows[0].cells[7].merge(t.rows[0].cells[8])
    # 合并 合计 (cols 9-10)
    t.rows[0].cells[9].merge(t.rows[0].cells[10])

    # 填充第1行（副表头）
    for i, h in enumerate(headers_sub):
        if h:
            cell = t.rows[1].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            r.font.size = Pt(10)
            r.font.name = '黑体'
            r.bold = True
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 填充数据行
    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 2].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str))
            r.font.size = Pt(10)
           .font.name = '仿宋'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 设置列宽
    col_widths = [Cm(1.5), Cm(1.8), Cm(1.8), Cm(1.5), Cm(1.5),
                  Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5),
                  Cm(1.8), Cm(1.8)]
    for i, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[i].width = w

    # 表后间距
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    pf = spacer.paragraph_format
    pf.line_spacing = Pt(4)
    pf.line_spacing_rule = docx_enum.WD_LINE_SPACING.EXACT


# ===== 构建正文 =====

# 标题
ap('鸿运承物业管理中心2026年一季度', 22, '宋体', c=True, ind=False)
ap('"接诉即办"工作分析讲评', 22, '宋体', c=True, ind=False)
ap('', ind=False)

# ---- 第一部分 ----
ap一、"接诉即办"派单分析讲评', 16, '黑体', b=True, ind=False)
ap('1. 市级派单情况', 16, '楷体', ind=False)
ap('2026年一季度，共签收市级派单246件；与2025年同期对比如下：')

# 创建复杂对比表
add_complex_table()

# 后续内容...
ap('一季度市级派单共246件，其中1月91件（诉求67件、表扬7件、其他17件），2月86件（诉求66件、表扬5件、其他15件），3月69件（诉求49件、表扬1件、其他18件、咨询1件）。诉求类派单182件，占派单总量的73.98%。')

ap('与2025年同期相比，2026年一季度市级派单总量246件，较2025年同期的251件减少5件，下降1.99%。从渠道分布看，2026年一季度所有派单均通过12345渠道接收，2025年同期除12345渠道的235件外，还有12328渠道11件、政风行风渠道件、舆情渠道4件。派单构成看，诉求类182件，较2025年同期的195件减少13件，下降6.67%；表扬类13件，较同期的10件增加3件，上升30%；其他类50件，较同期的44件增加6件，上升13.64%；咨询类1件，较同期的2件减少1件。')

ap('从诉求派单分布上看，诉求排名靠前的部室分别为：物管行政部78件（占42.86%），主要集中在车辆卫生问题；安保部43件（占23.63%），主要集中于场站安保问题、劳动人员管理及劳动纪律问题；人力资源部24件（占13.19%），集中在劳动人员调配及薪酬问题；工程设备部18件（占9.89%），集中在站台护栏问题、站务设施设置问题；餐饮服务部11件（占6.04%），集中在就餐服务问题。')

ap('与2025年同期相比，各专业部室诉求量变化如下：物管行政部78件，较2025年同期的62件增加16件，上升25.81%；安保部43件，较同期的33件增加10件，上升30.30%；人力资源部24件，较同期的34件减少10件，下降29.41%；工程设备部18件，较同期的30件减少12件，下降40.00%；餐饮服务部11件，较同期的13件减少2件，下降15.38%。物管行政部和安保部诉求量同比上升明显，人力资源部、工程设备部诉求量较去年同期有所下降。')

ap('总体来看，一季度市级派单的主要诉求来源仍集中在物管行政部及安保部。其中诉求排名靠前的问题项分别为车辆卫生问题、场站安保问题、劳动人员调配及薪酬问题、站台护栏问题、就餐服务问题及站务设施设置问题。')

ap('一季度诉求派单按业务/政策分类，业务类诉求148件（占81.32%），政策类诉求31件（占17.03%），纠纷类3件（占1.65%）。与2025年同期相比，业务诉求较同期的113件增加35件，上升30.97%；政策类诉求较同期的70件大幅下降39件，下降55.71%，反映出综A岗改革等政策经过前期宣贯和调整，相关争议已有所缓解。')

ap('从诉求派单按分中心分布看，承办量靠前的分中心分别为：动物园分中心27件、机关19件、西客站分中心16件、小营分中心14件、西苑分中心13件、天宫院分中心10件、方庄分中心9件、东直门分中心8件、南湖分中心8件、单店分中心7件、史各庄分中心7件。')

ap('2. 集团级派单情况', 16, '楷体', ind=False)
ap('2026年一季度，共接收"96166"派单220件（1月87件、2月84件、3月49件）。按性质分类：诉求102件、表扬16件、其他102件。')
ap('从部室分布看，"96166"派单主要承接部室为：物管行政部81件、安全保卫部（应急管理中心）40件、人力资源部37件、餐饮服务部17件、工程设备部14件。')

# ---- 第二部分 ----
ap('二、一季度"三率"考核情况分析', 16, '黑体', b=True, ind=False)
ap('2026年一季度，物业管理中心承办的市级"12345"派单中，纳入市中心回访考核共计94件。')
ap('从考核周期看，1月28件、2月37件、3月29件（截至3月12日）。')
ap('从评价情况看，一季度纳入考核派单全部为"只评价响应率"派单（94件，100%），无"三率"全考核派单。')
ap('从响应情况看，纳入考核的94件派单全部实现"已联系"，响应率为100%。')
ap('从解决率、满意率看，一季度无纳入解决率、满意率考核的派单，未出现失分情况。')
ap('与2025年上半年考核情况相比（221件纳入考核、其中2件失分），2026年一季度未出现失分情况，办理质量有所提升。但也应清醒认识到，仅考核响应率的格局下，后续纳入全口径考核的派单量可能增加，各单位仍需高度重视办理质量。')

# ---- 第三部分 ----
ap('三、突出诉求问题分析', 16, '黑体', b=True, ind=False)
ap('1. 车辆卫生问题（一季度高频问题）', 16, '楷体', ind=False)
ap('一季度车辆卫生问题是物业管理的重点和难点，涉及物管行政部多个分中心。反映较为集中的问题包括：车内座椅有灰尘、垃圾未清理、扶手脏粘、玻璃脏、车内异味等。涉及的线路较多，如107路、120路、362路、575路、599路、883路、896路、142路、387路等。')
ap('从办理情况看，各分中心均能及时响应，但存在反复投诉现象，日常保洁质量仍需持续加强监督和长效管理。')

ap('2. 场站安保问题', 16, '楷体', ind=False)
ap('一季度场站安保类诉求涉及安保部多个分中心，主要包括：保安工作态度、门岗管理、外卖进入场站管理、私家车违规停放及停车收费等。涉及小营、方庄、单店、南湖、田顺庄、辛庄等分中心。')

ap('3. 综合站务员A岗改革相关诉求', 16, '楷体', ind=False)
ap('一季度涉及人力资源部的政策类诉求较为集中，主要为综合站务员A岗优化配置改革引发的劳动人员调配及薪酬问题。涉及凌家庙、九龙、单店等多个分中心。')

ap('4. 就餐服务问题', 16, '楷体', ind=False)
ap('涉及餐饮服务部的诉求主要集中在食堂管理方面，包括菜品种类、价格合理性、工作人员态度等问题。')

ap('5. 站台护栏及站务设施问题', 16, '楷体', ind=False)
ap('涉及工程设备部的主要在站台护栏维修调整、站务设施设置等方面。')

# ---- 第四部分 ----
ap('四、下一步工作提示', 16, '黑体', b=True, ind=False)
ap('1. 深化"未诉先办"改革，推动主动治理', 16, '楷体', ind=False)
ap('依据集团公司《关于进一步深化接诉即办改革的意见》及物业管理中心修订后的实施细则，各单位要从"接诉即办"向"未诉先办"深化，加强日常巡查和问题预判，提前制定预防措施。')

ap('2. 严格落实分级协同和提级办理机制，提升处置能力', 16, '楷体', ind=False)
ap('对于基层分中心难以解决的难点问题，相关专业部门要主动提级办理、提级答复、提级见面。对于重复诉求，要及时与诉求人见面沟通。')

ap('3. 加强内部职工诉求的化解工作', 16, '楷体', ind=False)
ap('集团公司1至3月内部职工诉求持续较高（1月26件、2月33件、3月25件），主要集中在人力资源政策、薪资福利等方面。各相关专业及分中心要主动开展沟通工作，及时化解内部矛盾。')

ap('4. 规范剔除材料与挂账管理，降低被考核风险', 16, '楷体', ind=False)
ap('各专业部室和分中心要加强剔除材料的规范化管理。对符合"五挂两不挂"标准的合理诉求按时申报挂账，确保按期销账出库。')

ap('5. 高度重视"民生类"诉求办理质量', 16, '楷体', ind=False)
ap('物管行政部及安保部需重点加强车辆卫生及场站安保等民生类诉求的办理质量。对涉及"七有五性"指标的派单优先办理、重点督办。')

# ---- 第五部分 ----
ap('五、2026年一季度各分中心诉求承办量通报', 16, '黑体', b=True, ind=False)
ap('一季度各分中心诉求承办量如下：动物园分中心27件、机关19件、西客站分中心16件、小营分中心14件、西苑分中心13件、天宫院分中心10件、方庄分中心9件、东直门分中心8件、南湖分中心8件、单店分中心7件、史各庄分中心7件、凌家庙分中心6件、回龙观分中心5件、田顺庄分中心4件、大观园分中心4件、河滩分中心3件、明月湾分中心3件、四惠分中心3件、城南嘉园分中心3件、阜成门分中心3件、土桥分中心3件、九龙分中心2件、辛庄分中心2件、龙之乡分中心2件、小区留守组1件。')

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
    '鸿运承物业管理中心2026年一季度"接诉即办"工作分析讲评.docx')
doc.save(output_path)
print(f'✅ 文档已成功生成: {output_path}')
